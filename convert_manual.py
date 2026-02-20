"""HTMLマニュアルからDOCX・MD形式を生成するスクリプト"""
import re
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).parent
HTML_FILE = BASE_DIR / "ユーザマニュアル_手術室ガントチャート.html"
DOCX_FILE = BASE_DIR / "ユーザマニュアル_手術室ガントチャート.docx"
MD_FILE = BASE_DIR / "ユーザマニュアル_手術室ガントチャート.md"


def get_text(el):
    """要素からテキストを取得（code要素は前後にバッククォートなし）"""
    return el.get_text()


def html_to_md(html_content):
    """HTMLをMarkdownに変換"""
    soup = BeautifulSoup(html_content, "html.parser")
    container = soup.find("div", class_="container")
    if not container:
        container = soup.body or soup

    lines = []

    def process_inline(el):
        """インライン要素をMarkdownテキストに変換"""
        if isinstance(el, NavigableString):
            return str(el)
        tag = el.name
        if tag == "br":
            return "  \n"
        if tag == "code":
            return f"`{el.get_text()}`"
        if tag in ("strong", "b"):
            inner = "".join(process_inline(c) for c in el.children)
            return f"**{inner}**"
        if tag in ("em", "i"):
            inner = "".join(process_inline(c) for c in el.children)
            return f"*{inner}*"
        if tag == "a":
            inner = "".join(process_inline(c) for c in el.children)
            href = el.get("href", "")
            return f"[{inner}]({href})"
        if tag == "span":
            return "".join(process_inline(c) for c in el.children)
        return el.get_text()

    def inline_text(el):
        return "".join(process_inline(c) for c in el.children).strip()

    for el in container.children:
        if isinstance(el, NavigableString):
            t = str(el).strip()
            if t:
                lines.append(t)
            continue

        tag = el.name

        if tag == "h1":
            lines.append(f"# {el.get_text(separator=' ').strip()}")
            lines.append("")

        elif tag == "h2":
            lines.append(f"## {el.get_text().strip()}")
            lines.append("")

        elif tag == "h3":
            lines.append(f"### {el.get_text().strip()}")
            lines.append("")

        elif tag == "h4":
            lines.append(f"#### {el.get_text().strip()}")
            lines.append("")

        elif tag == "div" and "header-info" in el.get("class", []):
            spans = el.find_all("span")
            lines.append(" | ".join(s.get_text().strip() for s in spans))
            lines.append("")

        elif tag == "div" and "toc" in el.get("class", []):
            lines.append("## 目次")
            lines.append("")
            for li in el.find_all("li"):
                a = li.find("a")
                if a:
                    lines.append(f"- [{a.get_text()}]({a.get('href', '')})")
            lines.append("")

        elif tag == "div" and ("note" in el.get("class", []) or
                               "warning" in el.get("class", []) or
                               "danger" in el.get("class", [])):
            cls = el.get("class", ["note"])[0]
            prefix = {"note": "INFO", "warning": "WARNING", "danger": "DANGER"}.get(cls, "NOTE")
            # Process child elements within the note/warning/danger block
            inner_lines = []
            for child in el.children:
                if isinstance(child, NavigableString):
                    t = str(child).strip()
                    if t:
                        inner_lines.append(t)
                elif child.name == "table":
                    inner_lines.append("")
                    inner_lines.extend(_table_to_md(child))
                    inner_lines.append("")
                elif child.name == "ul":
                    for li in child.find_all("li", recursive=False):
                        inner_lines.append(f"  - {inline_text(li)}")
                elif child.name == "h4":
                    inner_lines.append("")
                    inner_lines.append(f"#### {child.get_text().strip()}")
                    inner_lines.append("")
                elif child.name == "p":
                    inner_lines.append(inline_text(child))
                else:
                    t = inline_text(child) if hasattr(child, 'children') else str(child).strip()
                    if t:
                        inner_lines.append(t)
            block_text = "\n".join(inner_lines).strip()
            lines.append(f"> **[{prefix}]** {block_text}")
            lines.append("")

        elif tag == "div" and ("terminal" in el.get("class", []) or
                               "folder-tree" in el.get("class", [])):
            lines.append("```")
            lines.append(el.get_text().strip())
            lines.append("```")
            lines.append("")

        elif tag == "div" and "step-box" in el.get("class", []):
            for child in el.children:
                if isinstance(child, NavigableString):
                    continue
                if child.name == "h3":
                    lines.append(f"### {child.get_text().strip()}")
                    lines.append("")
                elif child.name == "p":
                    lines.append(inline_text(child))
                    lines.append("")

        elif tag == "table":
            lines.extend(_table_to_md(el))
            lines.append("")

        elif tag == "p":
            lines.append(inline_text(el))
            lines.append("")

    return "\n".join(lines)


def _table_to_md(table_el):
    """HTML tableをMarkdownテーブルに変換"""
    rows = table_el.find_all("tr")
    if not rows:
        return []

    md_lines = []
    for i, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        cell_texts = []
        for c in cells:
            text = c.get_text(separator=" ").strip().replace("|", "\\|")
            text = re.sub(r"\s+", " ", text)
            cell_texts.append(text)
        md_lines.append("| " + " | ".join(cell_texts) + " |")
        if i == 0:
            md_lines.append("| " + " | ".join("---" for _ in cell_texts) + " |")
    return md_lines


def html_to_docx(html_content):
    """HTMLをDOCXに変換"""
    soup = BeautifulSoup(html_content, "html.parser")
    container = soup.find("div", class_="container")
    if not container:
        container = soup.body or soup

    doc = Document()

    # スタイル設定
    style = doc.styles["Normal"]
    style.font.name = "Meiryo UI"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo UI")

    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Meiryo UI"
        hs.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo UI")

    def add_inline_runs(paragraph, el):
        """インライン要素をparagraphにrunとして追加"""
        for child in el.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    paragraph.add_run(text)
            elif child.name == "br":
                paragraph.add_run("\n")
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x99)
            elif child.name in ("strong", "b"):
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ("em", "i"):
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "span":
                add_inline_runs(paragraph, child)
            elif child.name == "a":
                run = paragraph.add_run(child.get_text())
                run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
            else:
                paragraph.add_run(child.get_text())

    def add_table(table_el):
        """HTML tableをDOCX tableとして追加"""
        rows = table_el.find_all("tr")
        if not rows:
            return
        first_cells = rows[0].find_all(["th", "td"])
        ncols = len(first_cells)
        tbl = doc.add_table(rows=0, cols=ncols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            doc_row = tbl.add_row()
            for j, cell in enumerate(cells):
                if j < ncols:
                    doc_cell = doc_row.cells[j]
                    doc_cell.text = ""
                    p = doc_cell.paragraphs[0]
                    p.style = doc.styles["Normal"]
                    text = cell.get_text(separator=" ").strip()
                    text = re.sub(r"\s+", " ", text)
                    run = p.add_run(text)
                    run.font.size = Pt(9)
                    if cell.name == "th":
                        run.bold = True
                        # ヘッダ背景色
                        shading = doc_cell._element.get_or_add_tcPr()
                        shading_el = shading.makeelement(qn("w:shd"), {
                            qn("w:fill"): "2980B9",
                            qn("w:val"): "clear",
                        })
                        shading.append(shading_el)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def add_note_block(el, prefix="INFO", color=RGBColor(0x29, 0x80, 0xB9)):
        """注記ブロックを追加"""
        p = doc.add_paragraph()
        run = p.add_run(f"[{prefix}] ")
        run.bold = True
        run.font.color.rgb = color
        # テキスト部分を追加
        text = el.get_text(separator=" ").strip()
        text = re.sub(r"\s+", " ", text)
        p.add_run(text)
        p.paragraph_format.left_indent = Cm(0.5)

    for el in container.children:
        if isinstance(el, NavigableString):
            t = str(el).strip()
            if t:
                doc.add_paragraph(t)
            continue

        tag = el.name

        if tag == "h1":
            text = el.get_text(separator=" ").strip()
            doc.add_heading(text, level=1)

        elif tag == "h2":
            doc.add_heading(el.get_text().strip(), level=2)

        elif tag == "h3":
            doc.add_heading(el.get_text().strip(), level=3)

        elif tag == "h4":
            doc.add_heading(el.get_text().strip(), level=4)

        elif tag == "div" and "header-info" in el.get("class", []):
            spans = el.find_all("span")
            p = doc.add_paragraph()
            p.add_run(" | ".join(s.get_text().strip() for s in spans))
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        elif tag == "div" and "toc" in el.get("class", []):
            doc.add_heading("目次", level=2)
            for li in el.find_all("li"):
                doc.add_paragraph(li.get_text().strip(), style="List Bullet")

        elif tag == "div" and "note" in el.get("class", []):
            add_note_block(el, "INFO", RGBColor(0x29, 0x80, 0xB9))

        elif tag == "div" and "warning" in el.get("class", []):
            add_note_block(el, "WARNING", RGBColor(0xE6, 0x7E, 0x22))

        elif tag == "div" and "danger" in el.get("class", []):
            add_note_block(el, "DANGER", RGBColor(0xE7, 0x4C, 0x3C))

        elif tag == "div" and ("terminal" in el.get("class", []) or
                               "folder-tree" in el.get("class", [])):
            p = doc.add_paragraph()
            run = p.add_run(el.get_text().strip())
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Cm(1)

        elif tag == "div" and "step-box" in el.get("class", []):
            for child in el.children:
                if isinstance(child, NavigableString):
                    continue
                if child.name == "h3":
                    doc.add_heading(child.get_text().strip(), level=3)
                elif child.name == "p":
                    p = doc.add_paragraph()
                    add_inline_runs(p, child)

        elif tag == "table":
            add_table(el)
            doc.add_paragraph()  # 表後にスペース

        elif tag == "p":
            p = doc.add_paragraph()
            add_inline_runs(p, el)

        elif tag == "div" and "footer" in el.get("class", []):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(el.get_text().strip())
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(str(DOCX_FILE))
    print(f"DOCX生成完了: {DOCX_FILE}")


def main():
    html_content = HTML_FILE.read_text(encoding="utf-8")

    # DOCX生成
    html_to_docx(html_content)

    # MD生成
    md_content = html_to_md(html_content)
    MD_FILE.write_text(md_content, encoding="utf-8")
    print(f"MD生成完了: {MD_FILE}")


if __name__ == "__main__":
    main()
